#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
gen_cover_letter_docx.py — cover_letter_draft.md → cover_letter.docx

用法：
  python scripts/gen_cover_letter_docx.py --job_folder <job_folder> [--uid leon]

示例：
  python scripts/gen_cover_letter_docx.py --job_folder group-pmo_Oscar-Bravo_Junior-PMO_20260528 --uid leon
"""

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Mm, RGBColor
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Reuse the MD parser from the PDF script
sys.path.insert(0, str(Path(__file__).resolve().parent))
from gen_cover_letter_pdf import parse_cover_letter

PROJECT_DIR = Path(__file__).resolve().parent.parent
USERS_DIR   = PROJECT_DIR / "users"

MARGIN_MM = 22
FONT_NAME = "Times New Roman"


def _set_font(run, bold=False, size_pt=10.5, color=None):
    run.font.name = FONT_NAME
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)


def _set_para_spacing(para, before_pt=0, after_pt=6, line_rule=WD_LINE_SPACING.SINGLE):
    fmt = para.paragraph_format
    fmt.space_before = Pt(before_pt)
    fmt.space_after  = Pt(after_pt)
    fmt.line_spacing_rule = line_rule


def _add_horizontal_rule(doc):
    """Add a thin grey bottom border to a blank paragraph to simulate <hr>."""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(2)
    para.paragraph_format.space_after  = Pt(4)

    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")        # 0.5pt line
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B4B4B4")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(md_path: Path, docx_path: Path) -> None:
    md_text  = md_path.read_text(encoding="utf-8")
    sections = parse_cover_letter(md_text)

    doc = Document()

    # ── Page setup: A4, 22mm margins ──────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Mm(210)
    section.page_height = Mm(297)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(section, attr, Mm(MARGIN_MM))

    # Remove default empty paragraph added by python-docx
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # ── Header block ───────────────────────────────────────────────────────
    header_lines = sections["header"]
    if header_lines:
        # First line = sender name
        name_para = doc.add_paragraph()
        name_run  = name_para.add_run(header_lines[0])
        _set_font(name_run, bold=True, size_pt=13)
        _set_para_spacing(name_para, before_pt=0, after_pt=2)

        # Remaining header lines (contact info, date, company/city)
        for line in header_lines[1:]:
            p   = doc.add_paragraph()
            run = p.add_run(line)
            _set_font(run, size_pt=9, color=(80, 80, 80))
            _set_para_spacing(p, before_pt=0, after_pt=1)

    _add_horizontal_rule(doc)

    # ── Body paragraphs ────────────────────────────────────────────────────
    for para_text in sections["paragraphs"]:
        p   = doc.add_paragraph()
        run = p.add_run(para_text)
        _set_font(run, size_pt=10.5, color=(40, 40, 40))
        _set_para_spacing(p, before_pt=0, after_pt=6)

    # ── Closing signature ──────────────────────────────────────────────────
    if sections["closing"]:
        doc.add_paragraph()   # blank line before signature
        p   = doc.add_paragraph()
        run = p.add_run(sections["closing"])
        _set_font(run, size_pt=10.5)
        _set_para_spacing(p, before_pt=4, after_pt=0)

    doc.save(str(docx_path))
    print(f"DOCX generated: {docx_path}  ({docx_path.stat().st_size // 1024} KB)")


def main():
    parser = argparse.ArgumentParser(description="Convert cover_letter_draft.md → cover_letter.docx")
    parser.add_argument("--job_folder", required=True, help="Job folder name under users/{uid}/output/")
    parser.add_argument("--uid", default="leon", help="User ID (default: leon)")
    args = parser.parse_args()

    output_dir = USERS_DIR / args.uid / "output"
    job_dir    = output_dir / args.job_folder
    docx_path  = job_dir / "cover_letter.docx"

    # Resolve MD source (draft preferred, legacy fallback)
    md_path = job_dir / "cover_letter_draft.md"
    if not md_path.exists():
        alt = job_dir / "cover_letter.md"
        if alt.exists():
            md_path = alt
        else:
            print(f"ERROR: cover letter not found at {md_path}", file=sys.stderr)
            sys.exit(1)

    generate_docx(md_path, docx_path)


if __name__ == "__main__":
    main()
